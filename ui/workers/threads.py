import os
import shutil
import datetime
import fnmatch
from PyQt6.QtCore import QThread, pyqtSignal, QDate

_CHUNK = 4 * 1024 * 1024  # 4 MB per chunk


class FileOperationThread(QThread):
    """Thread for handling file copy/move operations without blocking the UI."""
    progress = pyqtSignal(int, str)
    finished_signal = pyqtSignal(bool, str)
    cancelled = pyqtSignal()

    def __init__(self, op_pairs: list, mode: str):
        super().__init__()
        self.op_pairs = op_pairs  # List of (src, target)
        self.mode = mode
        self.is_running = True
        self._total_bytes = 1
        self._done_bytes = 0

    def stop(self) -> None:
        self.is_running = False

    # ------------------------------------------------------------------
    # Internal helpers
    # ------------------------------------------------------------------

    def _scan_bytes(self, op_pairs: list) -> int:
        """Pre-scan total bytes so progress % is byte-accurate."""
        total = 0
        for src, _ in op_pairs:
            if os.path.isfile(src):
                try:
                    total += os.path.getsize(src)
                except OSError:
                    pass
            elif os.path.isdir(src):
                for dirpath, _, files in os.walk(src):
                    for f in files:
                        try:
                            total += os.path.getsize(os.path.join(dirpath, f))
                        except OSError:
                            pass
        return total or 1  # avoid division by zero

    def _chunked_copy(self, src: str, dst: str) -> bool:
        """Copy one file in chunks; returns False if cancelled mid-way."""
        dst_dir = os.path.dirname(dst)
        if dst_dir:
            os.makedirs(dst_dir, exist_ok=True)
        with open(src, "rb") as fin, open(dst, "wb") as fout:
            while True:
                if not self.is_running:
                    return False
                chunk = fin.read(_CHUNK)
                if not chunk:
                    break
                fout.write(chunk)
                self._done_bytes += len(chunk)
                pct = int(self._done_bytes / self._total_bytes * 100)
                self.progress.emit(pct, f"正在複製: {os.path.basename(src)}")
        shutil.copystat(src, dst)
        return True

    def _copy_tree(self, src_dir: str, dst_dir: str) -> bool:
        """Recursively copy a directory; returns False if cancelled."""
        os.makedirs(dst_dir, exist_ok=True)
        for entry in os.scandir(src_dir):
            if not self.is_running:
                return False
            dst_entry = os.path.join(dst_dir, entry.name)
            if entry.is_dir(follow_symlinks=False):
                if not self._copy_tree(entry.path, dst_entry):
                    return False
            else:
                if not self._chunked_copy(entry.path, dst_entry):
                    return False
        return True

    def _extract_archive(self, archive_path: str, dest_folder: str) -> None:
        """Extract archive to dest_folder with Smart Path logic, then send archive to recycle bin."""
        import zipfile, tarfile
        
        # Smart Path Logic: Decide whether to use the pre-created subfolder or parent folder
        parent_dir = os.path.dirname(dest_folder)
        final_dest = dest_folder
        
        _MAX_FILES = 10_000
        _MAX_BYTES = 2 * 1024 ** 3  # 2 GB

        try:
            if zipfile.is_zipfile(archive_path):
                with zipfile.ZipFile(archive_path, "r") as zf:
                    members = zf.infolist()
                    if not members: return
                    
                    # 智慧判定：檢查是否所有檔案都在同一個根目錄下
                    top_levels = set()
                    for m in members:
                        parts = m.filename.split('/')
                        if parts[0]: top_levels.add(parts[0])
                    
                    # 如果只有一個根目錄項目，且 ZIP 不是空包彈，則直接解壓到父目錄（不額外套娃）
                    if len(top_levels) == 1:
                        final_dest = parent_dir
                    
                    total_size = sum(m.file_size for m in members)
                    if len(members) > _MAX_FILES or total_size > _MAX_BYTES:
                        raise OSError(f"壓縮包過大（{len(members)} 個檔案），請手動解壓")
                    
                    os.makedirs(final_dest, exist_ok=True)
                    total = len(members) or 1
                    for i, member in enumerate(members):
                        if not self.is_running: return
                        zf.extract(member, final_dest)
                        self.progress.emit(int((i+1)/total*100), f"解壓: {member.filename}")

            elif tarfile.is_tarfile(archive_path):
                with tarfile.open(archive_path, "r:*") as tf:
                    members = tf.getmembers()
                    if not members: return
                    
                    top_levels = set()
                    for m in members:
                        parts = m.name.split('/')
                        if parts[0]: top_levels.add(parts[0])
                        
                    if len(top_levels) == 1:
                        final_dest = parent_dir
                        
                    total_size = sum(m.size for m in members)
                    if len(members) > _MAX_FILES or total_size > _MAX_BYTES:
                        raise OSError(f"壓縮包過大（{len(members)} 個檔案），請手動解壓縮")

                    os.makedirs(final_dest, exist_ok=True)
                    total = len(members) or 1
                    for i, member in enumerate(members):
                        if not self.is_running: return
                        tf.extract(member, final_dest)
                        self.progress.emit(int((i+1)/total*100), f"解壓: {member.name}")
            else:
                raise OSError(f"不支援的壓縮格式: {os.path.basename(archive_path)}")

            # 過河拆橋 (Trash after success)
            try:
                import send2trash as _s2t
                _s2t.send2trash(os.path.normpath(archive_path))
            except Exception:
                pass
        except Exception as e:
            raise e

    def _zip_folder(self, folder_path: str, zip_path: str) -> None:
        """Zip folder_path → zip_path with per-file progress. Cleans up on cancel."""
        import zipfile

        all_files: list[str] = []
        for dirpath, _, files in os.walk(folder_path):
            for f in files:
                all_files.append(os.path.join(dirpath, f))
        total = len(all_files) or 1
        parent_of_folder = os.path.dirname(folder_path)

        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED, compresslevel=6) as zf:
            for i, abs_path in enumerate(all_files):
                if not self.is_running:
                    zf.close()
                    try:
                        os.remove(zip_path)
                    except OSError:
                        pass
                    return
                arc_name = os.path.relpath(abs_path, parent_of_folder)
                zf.write(abs_path, arc_name)
                self.progress.emit(
                    int((i + 1) / total * 100),
                    f"壓縮: {os.path.basename(abs_path)}",
                )

    def _hash_file(self, path: str) -> str:
        """Compute SHA-256 of a file; returns hex digest."""
        import hashlib
        h = hashlib.sha256()
        with open(path, "rb") as f:
            for chunk in iter(lambda: f.read(_CHUNK), b""):
                h.update(chunk)
        return h.hexdigest()

    def _verify_and_trash(self, src: str, dst: str) -> None:
        """Verify src→dst integrity then send src to recycle bin.
        Rolls back dst and raises OSError on checksum mismatch."""
        if os.path.isfile(src):
            self.progress.emit(
                int(self._done_bytes / self._total_bytes * 100),
                f"正在校驗: {os.path.basename(src)}",
            )
            if self._hash_file(src) != self._hash_file(dst):
                try:
                    os.remove(dst)
                except OSError:
                    pass
                raise OSError(f"校驗失敗，來源已保留: {os.path.basename(src)}")
        else:
            # Directory: verify every file
            for dirpath, _, files in os.walk(src):
                rel = os.path.relpath(dirpath, src)
                dst_dir_path = dst if rel == "." else os.path.join(dst, rel)
                for fname in files:
                    self.progress.emit(
                        int(self._done_bytes / self._total_bytes * 100),
                        f"正在校驗: {fname}",
                    )
                    src_f = os.path.join(dirpath, fname)
                    dst_f = os.path.join(dst_dir_path, fname)
                    if not os.path.exists(dst_f) or self._hash_file(src_f) != self._hash_file(dst_f):
                        try:
                            shutil.rmtree(dst)
                        except OSError:
                            pass
                        raise OSError(f"校驗失敗，來源已保留: {fname}")

        # Checksums match — send source to recycle bin
        try:
            import send2trash as _s2t
            _s2t.send2trash(os.path.normpath(src))
        except Exception:
            # Fallback: permanent delete if send2trash unavailable
            if os.path.isfile(src):
                os.remove(src)
            else:
                shutil.rmtree(src)

    # ------------------------------------------------------------------

    def run(self) -> None:
        # --- extract / zip: separate dispatch, no byte-scan needed ---
        if self.mode in ("extract", "zip"):
            for src, dst in self.op_pairs:
                if not self.is_running:
                    break
                try:
                    src = os.path.normpath(src)
                    dst = os.path.normpath(dst)
                    if self.mode == "extract":
                        self._extract_archive(src, dst)
                    else:
                        self._zip_folder(src, dst)
                except Exception as e:
                    self.finished_signal.emit(False, str(e))
                    return
            self.finished_signal.emit(True, "操作完成")
            return

        # --- copy / move ---
        self._total_bytes = self._scan_bytes(self.op_pairs)
        self._done_bytes = 0

        for src, target in self.op_pairs:
            if not self.is_running:
                break
            try:
                src = os.path.normpath(src)
                target = os.path.normpath(target)

                if self.mode == "copy":
                    if os.path.isdir(src):
                        ok = self._copy_tree(src, target)
                    else:
                        ok = self._chunked_copy(src, target)
                    if not ok:
                        self.cancelled.emit()
                        return
                else:  # move
                    # Same-drive: instant rename, no verification needed
                    try:
                        os.rename(src, target)
                        try:
                            self._done_bytes += os.path.getsize(target)
                        except OSError:
                            pass
                        pct = int(self._done_bytes / self._total_bytes * 100)
                        self.progress.emit(pct, f"正在移動: {os.path.basename(src)}")
                    except OSError:
                        # Cross-drive: chunked copy → checksum verify → send2trash
                        if os.path.isdir(src):
                            ok = self._copy_tree(src, target)
                        else:
                            ok = self._chunked_copy(src, target)
                        if not ok:
                            self.cancelled.emit()
                            return
                        self._verify_and_trash(src, target)
            except Exception as e:
                self.finished_signal.emit(False, str(e))
                return

        self.finished_signal.emit(True, "操作完成")

def _find_context(path: str, keyword: str) -> str:
    """找出關鍵字在檔案中的第一個位置，返回簡短描述字串。"""
    ext = os.path.splitext(path)[1].lower()
    kw = keyword.lower()
    try:
        if ext == ".pdf":
            import fitz
            with fitz.open(path) as doc:
                for i, page in enumerate(doc, 1):
                    if kw in page.get_text().lower():
                        return f"第{i}頁"
        elif ext == ".docx":
            import docx as _docx
            doc = _docx.Document(path)
            for i, p in enumerate(doc.paragraphs, 1):
                if kw in p.text.lower():
                    return f"第{i}段"
        elif ext in (".xlsx", ".xls"):
            import openpyxl
            wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
            for ws in wb.worksheets:
                for row in ws.iter_rows():
                    for cell in row:
                        if cell.value and kw in str(cell.value).lower():
                            return f"{ws.title}!{cell.coordinate}"
        else:
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                for i, line in enumerate(f, 1):
                    if kw in line.lower():
                        snippet = line.strip()[:40]
                        return f"第{i}行: {snippet}"
    except Exception:
        pass
    return ""


def _extract_text(path: str) -> str:
    """從檔案提取純文字（小寫），供內容搜尋比對。
    支援 pdf / docx / xlsx / 純文字，其餘返回空字串。"""
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".pdf":
            import fitz
            with fitz.open(path) as doc:
                return "".join(page.get_text() for page in doc).lower()
        if ext == ".docx":
            import docx as _docx
            doc = _docx.Document(path)
            return "\n".join(p.text for p in doc.paragraphs).lower()
        if ext in (".xlsx", ".xls"):
            import openpyxl
            wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
            parts = []
            for ws in wb.worksheets:
                for row in ws.iter_rows(values_only=True):
                    parts.extend(str(c) for c in row if c is not None)
            return " ".join(parts).lower()
        # 純文字
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read().lower()
    except Exception:
        return ""


class SearchThread(QThread):
    """Thread for background file search."""
    match_found = pyqtSignal(str, float, float, str)  # path, mtime, size, context
    progress = pyqtSignal(str)
    finished_signal = pyqtSignal(int)

    def __init__(self, root_path, pattern, content_text, date_from, date_to, min_size, max_size):
        super().__init__()
        self.root_path = root_path
        self.pattern = pattern.lower()
        self.content_text = content_text.lower()
        self.date_from = date_from # QDate
        self.date_to = date_to # QDate
        self.min_size = min_size
        self.max_size = max_size if max_size > 0 else float('inf')
        self.is_running = True

    def stop(self):
        self.is_running = False

    def run(self):
        matches = 0
        for root, dirs, files in os.walk(self.root_path):
            if not self.is_running:
                break
                
            self.progress.emit(f"正在搜尋: {root}")
            
            for name in files + dirs:
                if not self.is_running:
                    break
                    
                full_path = os.path.join(root, name)
                
                # 名稱匹配邏輯 (Name matching logic)
                if self.pattern:
                    name_lower = name.lower()
                    if '*' in self.pattern or '?' in self.pattern:
                        # 萬用字元模式 (Wildcard mode)
                        if not fnmatch.fnmatch(name_lower, self.pattern):
                            continue
                    else:
                        # 包含子字串模式 (Substring/contains mode)
                        if self.pattern not in name_lower:
                            continue
                    
                try:
                    stats = os.stat(full_path)
                    if not (self.min_size <= stats.st_size <= self.max_size):
                        continue
                        
                    mtime = datetime.date.fromtimestamp(stats.st_mtime)
                    q_mtime = QDate(mtime.year, mtime.month, mtime.day)
                    
                    if self.date_from.isValid() and q_mtime < self.date_from:
                        continue
                    if self.date_to.isValid() and q_mtime > self.date_to:
                        continue
                        
                    if self.content_text and os.path.isfile(full_path):
                        if self.content_text not in _extract_text(full_path):
                            continue
                    
                    ctx = _find_context(full_path, self.content_text) if self.content_text else ""
                    self.match_found.emit(full_path, stats.st_mtime, stats.st_size, ctx)
                    matches += 1
                except:
                    continue
        self.finished_signal.emit(matches)

class CompareThread(QThread):
    """Thread for folder comparison."""
    progress = pyqtSignal(str)
    finished_signal = pyqtSignal(list, int)

    def __init__(self, left_path, right_path):
        super().__init__()
        self.left_path = left_path
        self.right_path = right_path
        self.is_running = True
        self.total_count = 0

    def stop(self):
        self.is_running = False

    def run(self):
        diffs = []
        self.total_count = 0
        try:
            self.compare_recursive("", diffs)
            self.finished_signal.emit(diffs, self.total_count)
        except Exception as e:
            self.progress.emit(f"錯誤: {str(e)}")
            self.finished_signal.emit(diffs, self.total_count)

    def compare_recursive(self, rel_path, diff_s):
        if not self.is_running: return
        
        lp = os.path.join(self.left_path, rel_path)
        rp = os.path.join(self.right_path, rel_path)
        
        try:
            left_names = set(os.listdir(lp)) if os.path.isdir(lp) else set()
            right_names = set(os.listdir(rp)) if os.path.isdir(rp) else set()
            all_names = left_names | right_names
            
            for name in all_names:
                if not self.is_running: return
                if name == ".sync_trash": continue
                self.total_count += 1
                sub_rel = os.path.join(rel_path, name)
                l_item = os.path.join(self.left_path, sub_rel)
                r_item = os.path.join(self.right_path, sub_rel)
                
                exists_l = os.path.exists(l_item)
                exists_r = os.path.exists(r_item)
                
                if exists_l and not exists_r:
                    count = self.count_all_children(l_item)
                    self.total_count += count
                    diff_s.append({'rel_path': sub_rel, 'status': 'L_ONLY', 'is_dir': os.path.isdir(l_item)})
                elif exists_r and not exists_l:
                    count = self.count_all_children(r_item)
                    self.total_count += count
                    diff_s.append({'rel_path': sub_rel, 'status': 'R_ONLY', 'is_dir': os.path.isdir(r_item)})
                else: 
                    if os.path.isdir(l_item) and os.path.isdir(r_item):
                        self.compare_recursive(sub_rel, diff_s)
                    elif os.path.isfile(l_item) and os.path.isfile(r_item):
                        ls = os.stat(l_item)
                        rs = os.stat(r_item)
                        if ls.st_size != rs.st_size or abs(ls.st_mtime - rs.st_mtime) > 2:
                            diff_s.append({'rel_path': sub_rel, 'status': 'DIFF', 'is_dir': False, 
                                          'l_newer': ls.st_mtime > rs.st_mtime})
                                          
            self.progress.emit(f"掃描中 ({self.total_count} 個項目): {rel_path if rel_path else '根目錄'}")
        except:
            pass

    def count_all_children(self, path):
        if not os.path.isdir(path):
            return 0
        count = 0
        try:
            for root, dirs, files in os.walk(path):
                if not self.is_running: break
                count += len(dirs) + len(files)
        except:
            pass
        return count
